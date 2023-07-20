import streamlit as st
import tkinter as tk
from tkinter import filedialog
# import fileinput
import emoji
from docx import Document
from docx.shared import Inches
import pandas as pd



####################  functions


# Função para validar se tem todas as informações necessárias.
def validate():
    if "sheet_path" in st.session_state and "model_path" in st.session_state and "folder" in st.session_state:
        funcao_botao()
    else:
        st.session_state.success = "FALTA"    


# função do botão que vai executar
def funcao_botao():
    parametro1 = st.session_state.sheet_path
    parametro2 = st.session_state.model_path
    gerar_contratos(parametro1, parametro2)
    st.write("FUNCIONOU")


# Função para executar o script principal
def gerar_contratos(caminho_xlsx, caminho_docx):
    
    # >>>>>> EXTRACT <<<<<<
    
    doc_carregado = Document(caminho_docx)
    df_contratos = pd.read_excel(caminho_xlsx)

    # >>>>>> TRANSFORM <<<<<<

    # pra cada linha vou criar um documento novo
    for index, row in df_contratos.iterrows():

        # criando um documento
        document = Document()

        # adicionando o logotipo no cabeçalho

        # Acesso ao cabeçalho
        cabecalho = document.sections[0].header

        # Cria um parágrafo no cabeçalho
        paragrafo = cabecalho.paragraphs[0]

        # Caminho da imagem
        if st.session_state.image_path != "no_image":
            imagem = st.session_state.image_path

            # Adiciona a imagem ao cabeçalho
            paragrafo.alignment = 1
            paragrafo.add_run().add_picture(imagem, width=Inches(2))

        
        
                # Adiciona uma margem abaixo da imagem
        paragrafo_vazio = cabecalho.add_paragraph()
        paragrafo_vazio.space_after = Inches(0.5)


        # adicionando a imagem no footer

        # Acesso ao footer
        footer = document.sections[0].footer

        # Cria um parágrafo no footer
        paragrafo_footer = footer.paragraphs[0]

        # Caminho da imagem
        if st.session_state.image_footer_path != "no_image":
            imagem_footer = st.session_state.image_footer_path

            # Adiciona a imagem_footer ao footer
            paragrafo_footer.alignment = 1
            paragrafo_footer.add_run().add_picture(imagem_footer, width=Inches(2))




        # montando cada parágrafo

        for paragrafo in doc_carregado.paragraphs:

            alinhamento = paragrafo.alignment

            # checando se o parágrafo tem alguma variável para ser substituída. 
            if "{{{" not in paragrafo.text:
                # se não, copia o mesmo conteúdo.
                document.add_paragraph(paragrafo.text).alignment = alinhamento
                
            else:
                # Se tem, substitui antes e depois add o parágrafo novo
                texto = paragrafo.text

                # Passo em cada coluna pra perguntar se tem essa variável no parágrafo
                for coluna in df_contratos.columns:
                    variavel = "{{{" + coluna + "}}}"

                    # substituo a variável pelo valor da linha naquela coluna
                    if variavel in texto:
                        texto = texto.replace(variavel, str(row[coluna]))

                # adiciono o parágrafo com o texto novo
                document.add_paragraph(texto).alignment = alinhamento

        # >>>>>> LOAD <<<<<<
             
        # criando os nomes dinâmicos para os arquivos finais
        if "docs_name" in st.session_state:
            nome_arquivo = st.session_state.folder + "/" + st.session_state.docs_name + "_" + row.iloc[0].replace("/","_") + ".docx"
        else:
            nome_arquivo = st.session_state.folder + "/" + row.iloc[0].replace("/","_") + ".docx"
        
        document.save(nome_arquivo)

        st.session_state.cont += 1

    st.session_state.success = "OK"
    st.balloons()
    st.session_state.hide_button = True
        
        


def main():

    # Set up tkinter
    root = tk.Tk()
    root.withdraw()

    # Make folder picker dialog appear on top of other windows
    root.wm_attributes('-topmost', 1)

    # Initializing cont
    if not "cont" in st.session_state: 
        st.session_state.cont = 0

    # Initializing success marker
    if not "success" in st.session_state: 
        st.session_state.success = ""

    # Initializing hide button marker
    if not "hide_button" in st.session_state:
        st.session_state.hide_button = False

    # Show logo
    logo = "https://ispn.org.br/site/wp-content/uploads/2021/04/logo_ISPN_2021.png"
    
    with st.columns(3)[1]:
        st.image(logo, output_format="PNG", width=220)

    # Title and subtitle
    st.markdown("<h1 style='text-align: center;'>Gerador de Documentos</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Para gerar documentos em massa, siga os passos a seguir:</p>", unsafe_allow_html=True)

    # Uncomment this to see session_state object
    # "st.session_state object", st.session_state
    

    # Empty space
    placeholder = st.empty()
    placeholder.text(" ")

    
    # Step 1 - Header image ----------------------------------------------
    st.markdown("<h5>1. Escolha uma imagem para o cabeçalho:</h5>", unsafe_allow_html=True)
    
    option_image = st.selectbox('',('Logo do ISPN', 'Carregar outra imagem', 'Continuar sem imagem no cabeçalho'))

    if option_image == "Logo do ISPN":
        st.session_state.image_path = "images/logo_ISPN_2021.png"
    elif option_image == "Carregar outra imagem":
        clicked_image_select = st.button('Escolha outra imagem')
        if clicked_image_select:
            st.session_state.image_path = filedialog.askopenfile(title="Selecione uma imagem", filetypes=[("Arquivos de Imagem", "*.jpg"), ("Arquivos de Imagem", "*.jpeg"), ("Arquivos de Imagem", "*.png")], master=root).name       
    elif option_image == "Continuar sem imagem no cabeçalho":
        st.session_state.image_path = "no_image"
    
    # Checking data in session_state and giving status
    if "image_path" in st.session_state:
        if st.session_state.image_path == "no_image":
            st.write(":white_check_mark: Não usar imagem no cabeçalho.")
        else:
            col1, col2 = st.columns([1, 3])
            col1.write(":white_check_mark: Imagem escolhida:")
            col2.image(st.session_state.image_path, width=150)
    else:
        st.session_state.image_path = logo

     
    # Step 2 - Load sheet ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>2. Carregue a tabela com as informações:</h5>", unsafe_allow_html=True)
    
    st.markdown("<span style='color:gray;'>Apenas são aceitas tabelas no formato .xlsx.</span>", unsafe_allow_html=True)
    st.markdown("<span style='color:gray;'>A primeira coluna da tabela será usada no nome dos arquivos gerados.</span>", unsafe_allow_html=True)
    
    # Sheet picker button
    clicked_sheet = st.button('Escolha a tabela')
    if clicked_sheet:
        st.session_state.sheet_path = filedialog.askopenfile(title="Selecione a tabela .xlsx", filetypes=[("Arquivos Excel", "*.xlsx")], master=root).name
        
    # Checking data in session_state and giving status
    if "sheet_path" in st.session_state:
        st.write(":white_check_mark: Tabela escolhida: ", st.session_state.sheet_path)        
        
        
    # Step 3 - Load model ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>3. Carregue o modelo do documento:</h5>", unsafe_allow_html=True)
    st.markdown("<span style='color:gray;'>Apenas são aceitos documentos no formato .docx.</span>", unsafe_allow_html=True)
    st.markdown("<span style='color:gray;'>As variárveis no modelo devem ter o mesmo nome das colunas da planilha, embaladas com chaves triplas {{{ }}}.</span>", unsafe_allow_html=True)
    
    # Model picker button
    clicked_model = st.button('Escolha o modelo')
    if clicked_model:
        st.session_state.model_path = filedialog.askopenfile(title="Selecione o modelo .docx", filetypes=[("Arquivos Word", "*.docx")], master=root).name
        
    # Checking data in session_state and giving status
    if "model_path" in st.session_state:
        st.write(":white_check_mark: Modelo escolhido: ", st.session_state.model_path)      
   
    
    # Step 4 - Choose destination ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>4. Escolha uma pasta para guardar o resultado:</h5>", unsafe_allow_html=True)

    # Folder picker button
    clicked_folder = st.button('Escolha a pasta')
    if clicked_folder:
        st.session_state.folder = filedialog.askdirectory(title="Escolha uma pasta para salvar os documentos gerados", master=root)

    # Checking data in session_state and giving status
    if "folder" in st.session_state:
        st.write(":white_check_mark: Pasta escolhida para salvar os documentos: ", st.session_state.folder)  

    # Step 5 - Footer image ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>5. Escolha uma imagem para o rodapé:</h5>", unsafe_allow_html=True)
    
    option_image_footer = st.selectbox('',('Continuar sem imagem no rodapé', 'Carregar uma imagem'))

    if option_image_footer == "Continuar sem imagem no rodapé":
        st.session_state.image_footer_path = "no_image"
    elif option_image_footer == "Carregar uma imagem":
        clicked_footer_image_select = st.button('Escolha uma imagem')
        if clicked_footer_image_select:
            st.session_state.image_footer_path = filedialog.askopenfile(title="Selecione uma imagem para o rodapé", filetypes=[("Arquivos de Imagem", "*.jpg"), ("Arquivos de Imagem", "*.jpeg"), ("Arquivos de Imagem", "*.png")], master=root).name       

    # Checking data in session_state and giving status
    if "image_footer_path" in st.session_state:
        if st.session_state.image_footer_path == "no_image":
            st.write(":white_check_mark: Não usar imagem no rodapé.")
        else:
            col1, col2 = st.columns([1, 3])
            col1.write(":white_check_mark: Imagem escolhida:")
            col2.image(st.session_state.image_footer_path, width=150)
    
    # Step 6 - Choose files names ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>6. Qual nome deseja dar aos documentos?</h5>", unsafe_allow_html=True)
    
    #  Creating session_state variable
    # st.session_state.docs_name = ""

    # Checking data in session_state
    if "docs_name" in st.session_state:
        if st.session_state.docs_name == "":
            # st.write("É nulo")
            st.session_state.docs_name = st.text_input("Digite e aperte ENTER")
            
        else:
            # st.write("Tem valor")
            st.write(":white_check_mark: O nome dos arquivos será: ", st.session_state["docs_name"], " + ", "o texto da primeira coluna da tabela.")
    else:
        # st.write("Não tem nome")
        doc_name_input = st.text_input("Digite e aperte ENTER")
        if doc_name_input != "":
            st.session_state.docs_name = doc_name_input
       


    # Step 7 - Gererate docs ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>7. Clique no botão abaixo para gerar os documentos:</h5>", unsafe_allow_html=True)
    st.button("Gerar documentos !", type="primary", disabled=st.session_state.hide_button, on_click=validate)

    if st.button("Reiniciar"):
        st.session_state.clear()
        st.experimental_rerun()



    if st.session_state.success == "OK":
        st.success(str(st.session_state.cont) + ' documentos gerados com sucesso na pasta ' + st.session_state.folder)

    elif st.session_state.success == "FALTA":
        st.warning("Você precisa escolher a TABELA, o MODELO e a PASTA DE DESTINO (etapas 2, 3 e 4).")



if __name__ == "__main__":
    main()